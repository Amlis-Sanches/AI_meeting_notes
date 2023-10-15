#import section
import openai
import os
import sys
from pydub import AudioSegment
from docx import Document
from PyQt6.QtWidgets import QApplication, QInputDialog

#Pull Files section
#get your API key to run the code but save it outside of your git rapository. 

#---------------------------------------------Functions-------------------------------------------------------------

def get_api_key(api_key_file):
    with open(api_key_file, 'r') as file:
        api_key = file.read().strip()
    return api_key

#check what type the audio file is. 
def check_audio_format(file_path):
    _, extension = os.path.splitext(file_path)
    return extension.lower()

#Trim the audio file incase it is to long
def trim_audio(input_file, output_file, start_time, end_time):
    audio = AudioSegment.from_file(input_file)
    trimmed_audio = audio[start_time:end_time]
    trimmed_audio.export(output_file, format="wav")

#This is for the main part of the meeting notes code. Transcribing it into words will allow us to submit the text through the API key. This will send the audio file to whisper and then you will recieve text back.
def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = openai.Audio.transcribe("whisper-1",audio_file)
    return transcription['text']

#-----------------------Summarizing and analyzing the transcript with GPT-3------------------------------------------

#you can do this all in one function but it is found that splitting up the functions allow for higher quality. 
# Here we take the raw text and pass it through each function. 

def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

#--Summary extraction
#take the text and summarizes it into a concise abstract paragraph.
def abstract_summary_extraction(transcription):
    response = openai.Completion.create(
        model="text-davinci-003",
        temperature=0,
        prompt="You are a highly skilled AI trained in language comprehension and summarization. "
            "I would like you to read the following text and summarize it into a concise abstract paragraph. "
            "Aim to retain the most important points, providing a coherent and readable summary that could help a person "
            "understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary "
            "details or tangential points.\n\n" + transcription
    )
    return response['choices'][0]['text']



#--Extract Key Points
def key_points_extraction(transcription):
    response = openai.Completion.create(
        model="text-davinci-003",
        temperature=0,
        prompt="You are a proficient AI with a specialty in distilling information into key points." 
            "Based on the following text, identify and list the main points that were discussed or brought up." 
            "These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion." 
            "Your goal is to provide a list that someone could read to quickly understand what was talked about.\n\n" + transcription
    )
    return response['choices'][0]['text']

#--Action Item Extraction
def action_item_extraction(transcription):
    response = openai.Completion.create(
        model="text-davinci-003",
        temperature=0,
        prompt="You are an AI expert in analyzing conversations and extracting action items."
            "Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done." 
            "These could be tasks assigned to specific individuals, or general actions that the group has decided to take." 
            "Please list these action items clearly and concisely. \n\n" + transcription
    )
    return response['choices'][0]['text']

#--Sentiment Analysis
def sentiment_analysis(transcription):
    response = openai.Completion.create(
        model="text-davinci-003",
        temperature=0,
        prompt="As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text."
            "Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used." 
            "Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible.\n\n" + transcription
    )
    return response['choices'][0]['text']


#--Exporting Meeting Minutes
def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

def save_api_key(api_key):
    # Determine the current directory where the script is located
    current_directory = os.path.dirname(os.path.abspath(__file__))
    # Move up one directory to get the parent directory
    parent_directory = os.path.dirname(current_directory)
    # Create the path to the api_key.txt file within that directory
    api_key_file = os.path.join(parent_directory, "api_key.txt")
    # Save the API key to the generated path
    with open(api_key_file, 'w') as file:
        file.write(api_key)


def prompt_for_api_key():
    # Get user input for the API key
    api_key, ok = QInputDialog.getText(None, "API Key Input", "Please enter your API key:")

    # If the user pressed OK and provided an input, save it to a file
    if ok and api_key:
        save_api_key(api_key)


#----------------------------Main Body----------------------------------------
def main():

    audio_format = check_audio_format(audio_file_path)

    if audio_format != '.mp3':
        print(f'The file {audio_file_path} is not an MP3 file.')
        audio.export("EarningsCall.mp3", format="mp3")
        audio_file_path = "EarningsCall.mp3"


    # Get the duration in milliseconds
    duration_ms = len(audio_file)

    #reducing the size of the wav file
    output_file = "trimmed_audio.wav"
    start_time = 0  # Start time in milliseconds
    end_time = duration_ms/3   # End time in milliseconds

    input_file_path = r"C:\Users\natha\Documents\Coding\AI_meeting_notes\EarningsCall.wav"
    trim_audio(input_file_path, output_file, start_time, end_time)

    #use the Whisper model to take the audio and transcribe the file
    transcription = transcribe_audio(audio_file_path)
    print(transcription)
    minutes = meeting_minutes(transcription)
    print(minutes)

    save_as_docx(minutes, 'meeting_minutes.docx')

#----------------------------Interactive window for API key-----#
#This section asks for you to imput the API key for open AI and #
#store it on your hardrive main folder. This way the API key is #
#seporate from the program and you can remove when needed.      #
#---------------------------------------------------------------#

#check file path
current_directory = os.path.dirname(os.path.abspath(__file__))
# Move up one directory
parent_directory = os.path.dirname(current_directory)
api_key_file = os.path.join(parent_directory, "api_key.txt")# Path to the text file containing the API key

if os.path.exists(api_key_file):
    # Initialize the API with the retrieved key
    openai.api_key = get_api_key(api_key_file)
else:
    app = QApplication(sys.argv)
    prompt_for_api_key()
    app.exec()


#--------------------Prompting for Audio file-------------------#
#Here we are prompting for the audio file and asking it to be in#
#MP3. once a file has been selected it will check if its an mp3 #
#if its not it will ask for it again.                           #
#---------------------------------------------------------------#

# This is the main code, calling the main and aditional adjustment. 
# Load the audio file
audio_file = AudioSegment.from_file(r"C:\Users\natha\Documents\Coding\AI_meeting_notes\EarningsCall.wav")
# Load your WAV file
audio = AudioSegment.from_wav("EarningsCall.wav")

#if __name__ == "__main__":
#    main()